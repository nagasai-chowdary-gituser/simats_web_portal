from docx import Document
from docxcompose.composer import Composer

def remove_trailing_empty_paragraphs(doc):
    """
    Removes trailing empty paragraphs that cause blank pages.
    """
    while len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

def merge_docx(doc_paths, output_path):
    """
    Merge DOCX documents WITHOUT introducing blank pages.
    - Removes trailing empty paragraphs
    - Adds page breaks ONLY when required
    - Preserves all formatting
    """

    if not doc_paths:
        raise ValueError("No DOCX files provided for merging")

    # Load master document
    master = Document(doc_paths[0])
    composer = Composer(master)

    # Clean trailing empties in master
    remove_trailing_empty_paragraphs(master)

    for path in doc_paths[1:]:
        sub_doc = Document(path)

        # Clean trailing empty paragraphs in sub_doc
        remove_trailing_empty_paragraphs(sub_doc)

        # ⭐ CHECK: If master already ends with a page break, DO NOT add another
        last_paragraph = master.paragraphs[-1]
        last_run = last_paragraph.runs[-1] if last_paragraph.runs else None

        already_has_pagebreak = False
        if last_run and last_run._element.tag.endswith('br'):
            already_has_pagebreak = True

        # ⭐ Only add a page break when needed
        if not already_has_pagebreak:
            master.add_page_break()

        # Append doc
        composer.append(sub_doc)

    composer.save(output_path)
