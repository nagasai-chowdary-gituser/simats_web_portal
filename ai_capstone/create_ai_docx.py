import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
#   FIXED TABLE OF CONTENTS STRUCTURE (FINAL VERSION)
# ============================================================

TOC_STRUCTURE = [
    ("Abstract", ["—"], "6"),

    ("Chapter 1", [
        "1.1 Background Information",
        "1.2 Problem Definition",
        "1.3 Objectives of the Study",
        "1.4 Scope of the Project",
        "1.5 Methodology Overview"
    ], "7–11"),

    ("Chapter 2", [
        "2.1 Problem Analysis",
        "2.2 Evidence and Motivation",
        "2.3 Stakeholder Requirements",
        "2.4 Challenges Identified",
        "2.5 Supporting Research",
        "2.6 Summary"
    ], "12-16"),

    ("Chapter 3", [
        "3.1 System Design",
        "3.2 Architecture Overview",
        "3.3 Tools and Technologies",
        "3.4 Data Flow Diagram",
        "3.5 Functional Description",
        "3.6 Engineering Standards",
        "3.7 Design Justification",
        "3.8 Summary"
    ], "17-20"),

    ("Chapter 4", [
        "4.1 Implementation Details",
        "4.2 Algorithms Used",
        "4.3 Model or System Workflow",
        "4.4 Performance Evaluation",
        "4.5 Challenges Encountered",
        "4.6 Summary"
    ], "21-24"),

    ("Chapter 5", [
        "5.1 Learning Outcomes",
        "5.2 Collaboration and Teamwork",
        "5.3 Application of Standards",
        "5.4 Problem-solving Skills",
        "5.5 Industry Exposure",
        "5.6 Personal Growth",
        "5.7 Summary"
    ], "25-28"),

    ("Chapter 6", [
        "6.1 Summary of Work",
        "6.2 Major Findings",
        "6.3 Impact and Significance",
        "6.4 Limitations",
        "6.5 Future Scope",
        "6.6 Conclusion"
    ], "29-31"),

    ("References", ["—"], "32"),
    ("Appendices", ["—"], "33")
]

# ============================================================
#   STYLE FUNCTION
# ============================================================

def set_style(paragraph, *, bold=False, align="justify", size=12):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold

    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragraph.paragraph_format.line_spacing = 1.5

# ============================================================
#   CLEAN TEXT
# ============================================================

def clean_text(text: str):
    text = re.sub(r"#+ ", "", text)
    text = re.sub(r"\*", "", text)
    return text.strip()

# ============================================================
#   TABLE OF CONTENTS BUILDER  (DO NOT TOUCH)
# ============================================================

def insert_table_of_contents(doc):
    # Title
    title = doc.add_paragraph("TABLE OF CONTENTS")
    for r in title.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Table Grid
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["S.No.", "Chapter", "Topics", "Page No."]
    hdr = table.rows[0].cells

    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(8)
        run.font.name = "Times New Roman"
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sno = 1

    for chapter, topics, pages in TOC_STRUCTURE:
        row = table.add_row().cells

        # S.No
        p = row[0].paragraphs[0]
        r = p.add_run(str(sno))
        r.font.name = "Times New Roman"
        r.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Chapter Name
        p = row[1].paragraphs[0]
        r = p.add_run(chapter)
        r.font.name = "Times New Roman"
        r.font.size = Pt(8)
        r.bold = True

        # Topics
        p = row[2].paragraphs[0]
        r = p.add_run("\n".join(topics))
        r.font.name = "Times New Roman"
        r.font.size = Pt(8)

        # Page No
        p = row[3].paragraphs[0]
        r = p.add_run(pages)
        r.font.name = "Times New Roman"
        r.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sno += 1

# ============================================================
#   DOCX BUILDER  (FINAL BOSS VERSION)
# ============================================================

def create_ai_docx(sections: dict, output_path: str):
    doc = Document()

    # 1) Insert TOC
    insert_table_of_contents(doc)
    # Keep this page break so Abstract starts on next page
    doc.add_page_break()

    # 2) CHAPTERS IN FIXED ORDER
    for chapter_title, _, _ in TOC_STRUCTURE:

        # Find matching AI-generated key (e.g. "Chapter 1", "Abstract", "Appendices")
        real_key = None
        for key in sections.keys():
            if key.lower().startswith(chapter_title.lower()):
                real_key = key
                break

        if not real_key:
            continue

        content = sections[real_key]
        cleaned = clean_text(content)

        # ---------------------------------------------------
        # SPECIAL — REFERENCES
        # ---------------------------------------------------
        if chapter_title.lower() == "references":
            heading = doc.add_paragraph("REFERENCES")
            set_style(heading, bold=True, size=16, align="center")

            for line in cleaned.split("\n"):
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    set_style(p, bold=False, size=12, align="left")

            doc.add_page_break()
            continue

        # ---------------------------------------------------
        # CHAPTER / ABSTRACT / APPENDICES HEADING
        # ---------------------------------------------------
        ct = chapter_title.lower()

        if ct == "abstract":
            # Just "ABSTRACT"
            heading_text = "ABSTRACT"
        elif ct == "appendices":
            # Just "APPENDICES"
            heading_text = "APPENDICES"
        else:
            # Chapter 1, 2, ... (normal chapters)
            if ":" in real_key:
                heading_text = real_key.upper()
            else:
                first_line = cleaned.split("\n")[0].strip()
                if ":" in first_line:
                    heading_text = first_line.upper()
                else:
                    # fallback: CHAPTER 1  (no extra "INTRODUCTION" to avoid confusion)
                    heading_text = real_key.upper()

        heading = doc.add_paragraph()
        run = heading.add_run(heading_text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacing under heading

        # ---------------------------------------------------
        # CONTENT SECTION (FINAL BOSS VERSION)
        # ---------------------------------------------------
        for line in cleaned.split("\n"):
            line = line.strip()
            if not line:
                continue

            # Skip duplicate headings inside AI text
            if line.lower().startswith("chapter") or line.lower().startswith("abstract") or line.lower().startswith("appendices"):
                continue

            # 1) Handle BOLD markers from clean_auto_headings (Appendix A/B/C/D)
            if line.startswith("@@BOLD@@") and line.endswith("@@END@@"):
                pure = line.replace("@@BOLD@@", "").replace("@@END@@", "").strip()
                p = doc.add_paragraph(pure)
                set_style(p, bold=True, size=12, align="left")
                continue

            # 2) Backup: if AI outputs "Appendix A ..." without markers, still bold it
            if re.match(r"^Appendix\s+[A-Z]\b.*", line, re.IGNORECASE):
                p = doc.add_paragraph(line)
                set_style(p, bold=True, size=12, align="left")
                continue

            # 3) Bold numeric subheadings like "1.1 Background Information"
            if re.match(r"^\d+\.\d+\s", line):
                p = doc.add_paragraph(line)
                set_style(p, bold=True, size=12, align="left")
                continue

            # 4) Normal paragraph
            p = doc.add_paragraph(line)
            set_style(p, bold=False, size=12, align="justify")

        # Page break after each section
        doc.add_page_break()

    # SAVE FILE
    doc.save(output_path)
